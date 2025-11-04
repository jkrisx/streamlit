# app.py
import re, sqlite3, io, tempfile
from docx import Document
from openpyxl import load_workbook, Workbook
import streamlit as st
import os

st.set_page_config(page_title="AL Converter - Streamlit", layout="wide")

st.title("Analisis Lendutan Converter — Streamlit")
st.markdown("Upload file .docx (bisa banyak), file database .db dan template Excel (.xlsx). Klik **Run** untuk menghasilkan Excel.")

# --- Inputs ---
uploaded_docx = st.file_uploader("Upload file .docx (multiple)", type=["docx"], accept_multiple_files=True)
uploaded_db = st.file_uploader("Upload SQLite DB (.db)", type=["db", "sqlite", "sqlite3"])
uploaded_xlsx = st.file_uploader("Upload Excel template (.xlsx) (optional)", type=["xlsx"])
start_row = st.number_input("Start row in Excel (default 2)", min_value=1, value=2, step=1)

# Helper: save uploaded DB to temp file for sqlite3 to open
def save_tempfile(uploaded_file):
    tmp = tempfile.NamedTemporaryFile(delete=False)
    tmp.write(uploaded_file.read())
    tmp.flush()
    tmp.close()
    return tmp.name

# --- Core conversion logic (refactored & fixed) ---
class LendutanConverter:
    def __init__(self, docs, wb, db_path, start_row=2):
        self.docs = docs  # list of file-like objects
        self.wb = wb
        self.ws = self.wb.active
        self.conn = sqlite3.connect(db_path)
        self.cursor = self.conn.cursor()
        self.current_row = start_row
        self.processed_tables = set()

    def extract_voltase_from_text(self, text):
        m = re.search(r'\b(\d{3,4}kV)\b', text, re.IGNORECASE)
        return m.group(1) if m else None

    def extract_location_from_text(self, text):
        # find voltage location
        volt_m = re.search(r'\b\d{3,4}kV\b', text, re.IGNORECASE)
        if volt_m:
            start = volt_m.end()
            rest = text[start:]
        else:
            rest = text

        # if SPAN exists, cut before it
        span_m = re.search(r'\bSPAN\b', rest, re.IGNORECASE)
        if span_m:
            loc_raw = rest[:span_m.start()]
        else:
            loc_raw = rest

        loc_raw = loc_raw.strip()
        if not loc_raw:
            return None

        # Remove trailing numeric ranges like "13-22" or single numbers at the end
        loc_no_trailing_nums = re.sub(r'\s*\d+(?:[\-\&]\d+)?\s*$', '', loc_raw).strip()

        # Special-case mapping from your original script
        specials = {
            "JAWA 9&10 - CILEGON BARU": "JAWA9&10-CILEGONBARU",
            "SURALAYA - JAWA 9&10": "SURALAYA-JAWA9&10",
            "MUARA KARANG - DURIKOSAMBI": "MUARAKARANG-DURIKOSAMBI",
            "MUARA KARANG - DURI KOSAMBI": "MUARAKARANG-DURIKOSAMBI",
            "SURALAYA BARU - SURALAYA": "SURALAYA BARU - SURALAYA",
            "KEMBANGAN - DURIKOSAMBI": "KEMBANGAN - DURIKOSAMBI",
        }
        key = re.sub(r'\s+', ' ', loc_raw).strip().upper()
        if key in specials:
            return specials[key]

        # normalize: remove extra spaces, remove spaces inside tokens but keep '-' as separator
        parts = re.split(r'\s*-\s*', loc_no_trailing_nums)
        cleaned = []
        for p in parts:
            # collapse multiple spaces, then remove spaces within tokens (but keep 9&10 and similar)
            p2 = re.sub(r'\s+', ' ', p).strip().upper()
            # remove spaces between words (original 'Format C' seems to want no spaces within each part)
            p2 = p2.replace(' ', '')
            if p2:
                cleaned.append(p2)
        return "-".join(cleaned) if cleaned else None

    def get_title_text(self, doc: Document):
        # find paragraph containing kV
        for para in doc.paragraphs:
            if re.search(r'\b\d{3,4}kV\b', para.text, re.IGNORECASE):
                return para.text.strip()
        # fallback: join first three paragraphs
        return " ".join(p.text.strip() for p in doc.paragraphs[:3]).strip()

    def process_word_file(self, file_like, fname=None):
        try:
            doc = Document(file_like)
        except Exception as e:
            st.warning(f"File {fname or ''} cannot be opened as Word document. Skipped.")
            return

        title_text = self.get_title_text(doc)
        voltase = self.extract_voltase_from_text(title_text)
        lokasi = self.extract_location_from_text(title_text)

        if not voltase or not lokasi:
            st.info(f"Skip {fname or ''}: voltase or lokasi tidak terdeteksi.")
            return

        gabungan_value = f"{voltase} {lokasi}"

        if not doc.tables:
            st.info(f"Skip {fname or ''}: tidak ada tabel.")
            return

        table = doc.tables[0]
        table_data = tuple(tuple(cell.text.strip() for cell in row.cells) for row in table.rows)
        if table_data in self.processed_tables:
            st.info(f"Skip {fname or ''}: tabel duplikat.")
            return
        self.processed_tables.add(table_data)

        # build nomor_akhir from first column per row and query DB
        for row in table.rows[1:]:
            values = [cell.text.strip() for cell in row.cells]
            first_col = values[0] if values else ""
            match_span_num = re.search(r'\d+-(\d+)', first_col)
            nomor_akhir = match_span_num.group(1) if match_span_num else None
            nomor_akhir = nomor_akhir.zfill(4) if nomor_akhir else ""

            if lokasi.upper().replace(' ', '') == "KEMBANGAN-DURIKOSAMBI":
                # use functlog pattern with '-6' in original
                like_induk = f"%{gabungan_value}%"
                like_functlog = f"%-6{nomor_akhir}%"
                sql = "SELECT functlog FROM master_functlog WHERE induk LIKE ? AND functlog LIKE ? AND nama LIKE 'SPAN SUTET%' LIMIT 1"
                params = (like_induk, like_functlog)
            else:
                like_induk = f"%{gabungan_value}%"
                like_functlog = f"%-S%{nomor_akhir}%"
                sql = "SELECT functlog FROM master_functlog WHERE induk LIKE ? AND functlog LIKE ? AND nama LIKE 'SPAN SUTET%' LIMIT 1"
                params = (like_induk, like_functlog)

            try:
                self.cursor.execute(sql, params)
                res = self.cursor.fetchone()
                functlog_val = res[0] if res else ""
            except Exception as e:
                functlog_val = ""
                st.warning(f"DB query error for {gabungan_value}, row first col '{first_col}': {e}")

            # write to excel: functlog in col A, then table values starting col B (column index start=1)
            col_offset = 1
            self.ws.cell(row=self.current_row, column=1, value=functlog_val)
            for i, v in enumerate(values, start=2):
                self.ws.cell(row=self.current_row, column=i, value=v)
            self.current_row += 1

    def run_all(self):
        for file_like, fname in self.docs:
            st.write(f"🔍 Memproses file: {fname}")
            self.process_word_file(file_like, fname=fname)
        self.conn.close()
        st.success("✅ Proses selesai.")

# --- Run when user clicks ---
if st.button("Run Conversion"):
    if not uploaded_docx:
        st.error("Silakan upload minimal 1 file .docx.")
    elif not uploaded_db:
        st.error("Silakan upload file database (.db).")
    else:
        # prepare DB temp file
        db_path = save_tempfile(uploaded_db)

        # prepare workbook
        if uploaded_xlsx:
            try:
                xlsx_bytes = uploaded_xlsx.read()
                wb = load_workbook(io.BytesIO(xlsx_bytes))
            except Exception as e:
                st.error(f"Template Excel tidak bisa dibuka: {e}")
                wb = Workbook()
        else:
            wb = Workbook()

        # prepare docs as list of (file-like, filename)
        docs = []
        for f in uploaded_docx:
            # recreate BytesIO since Streamlit's uploaded file might be consumed
            f.seek(0)
            file_bytes = f.read()
            docs.append((io.BytesIO(file_bytes), f.name))

        converter = LendutanConverter(docs=docs, wb=wb, db_path=db_path, start_row=int(start_row))
        with st.spinner("Memproses..."):
            converter.run_all()

        # save workbook to bytes and offer download
        out = io.BytesIO()
        converter.wb.save(out)
        out.seek(0)
        st.download_button("Download Excel hasil konversi", data=out.getvalue(), file_name="analisis_lendutan_converted.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        # cleanup temp db
        try:
            os.unlink(db_path)
        except Exception:
            pass

st.markdown("---")
st.info("Catatan: Jika ingin agar aplikasi ini berjalan otomatis pada folder server (bukan via upload), deploy aplikasi pada server yang memiliki akses ke folder dan sesuaikan kode untuk membaca folder lokal.")
